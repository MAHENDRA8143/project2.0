from __future__ import annotations

import os
import re
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

SOURCE_PATH = Path(r"e:\prooooooooojjectt\docs\AI_Based_Smart_WWTP_Report_DOCX_Ready.md")
IMAGE_DIR = Path(r"e:\prooooooooojjectt\docs\gallery_images")
OUTPUT_PATH = Path(r"e:\prooooooooojjectt\docs\AI_Based_Smart_WWTP_Report_WITH_GALLERY_{0}.docx".format(datetime.now().strftime('%Y%m%d_%H%M%S')))

BODY_FONT = "Times New Roman"
BODY_SIZE = 12


def set_page_number_footer(doc: Document) -> None:
    for section in doc.sections:
        footer = section.footer
        paragraph = footer.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run()
        fld_char_begin = OxmlElement("w:fldChar")
        fld_char_begin.set(qn("w:fldCharType"), "begin")
        instr_text = OxmlElement("w:instrText")
        instr_text.set(qn("xml:space"), "preserve")
        instr_text.text = "PAGE"
        fld_char_end = OxmlElement("w:fldChar")
        fld_char_end.set(qn("w:fldCharType"), "end")
        run._r.append(fld_char_begin)
        run._r.append(instr_text)
        run._r.append(fld_char_end)
        for r in paragraph.runs:
            r.font.name = BODY_FONT
            r.font.size = Pt(10)


def configure_styles(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = Pt(BODY_SIZE)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)


def set_margins(doc: Document) -> None:
    for section in doc.sections:
        section.left_margin = Inches(1.5)
        section.right_margin = Inches(1.0)
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)


def format_run(run, *, size=BODY_SIZE, bold=False, italic=False, color=None) -> None:
    run.font.name = BODY_FONT
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color is not None:
        run.font.color.rgb = color
    run._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)


def add_paragraph(doc: Document, text: str, *, align=WD_ALIGN_PARAGRAPH.JUSTIFY, size=BODY_SIZE, bold=False, italic=False, space_after=6):
    p = doc.add_paragraph()
    p.alignment = align
    r = p.add_run(text)
    format_run(r, size=size, bold=bold, italic=italic)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(space_after)
    return p


def add_title(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    format_run(r, size=16, bold=True)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(8)


def add_section_heading(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(text)
    format_run(r, size=14, bold=True)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(4)


def add_subheading(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(text)
    format_run(r, size=12, bold=True)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)


def add_image(doc: Document, image_path: Path, width_inches: float = 6.5) -> None:
    if not image_path.exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(image_path), width=Inches(width_inches))
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.0


def add_page_label(doc: Document, page_num: int) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"=== Page {page_num} ===")
    format_run(r, size=12, bold=True, color=RGBColor(70, 70, 70))
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.5


def parse_markdown_sections(text: str) -> list[tuple[int, str, list[str]]]:
    pages: list[tuple[int, str, list[str]]] = []
    chunks = re.split(r"^=== Page (\d+) ===$", text, flags=re.M)
    if len(chunks) < 3:
        raise ValueError("Source markdown does not contain page markers.")

    # chunks format: [preface, page_num, page_body, page_num, page_body, ...]
    for i in range(1, len(chunks), 2):
        page_num = int(chunks[i])
        body = chunks[i + 1].strip()
        lines = [line.rstrip() for line in body.splitlines()]
        pages.append((page_num, body, lines))
    return pages


def render_page(doc: Document, page_num: int, body: str, lines: list[str]) -> None:
    add_page_label(doc, page_num)

    # Extract title from the first markdown heading if present.
    idx = 0
    while idx < len(lines) and not lines[idx].strip():
        idx += 1
    if idx < len(lines) and lines[idx].startswith("## "):
        add_title(doc, lines[idx][3:].strip())
        idx += 1

    while idx < len(lines):
        line = lines[idx].rstrip()
        if not line:
            idx += 1
            continue
        if line.startswith("[INSERT FIGURE"):
            add_subheading(doc, line.strip())
            figure_number_match = re.search(r"(\d+)", line)
            figure_number = int(figure_number_match.group(1)) if figure_number_match else None
            if idx + 3 < len(lines):
                add_paragraph(doc, lines[idx + 1], align=WD_ALIGN_PARAGRAPH.LEFT, italic=False)
                add_paragraph(doc, lines[idx + 2], align=WD_ALIGN_PARAGRAPH.LEFT)
                add_paragraph(doc, lines[idx + 3], align=WD_ALIGN_PARAGRAPH.LEFT)
                idx += 4
            else:
                idx += 1
            if figure_number is not None:
                image_path = IMAGE_DIR / f"panel-{figure_number:02d}.png"
                add_image(doc, image_path)
                if figure_number == 20 and not image_path.exists():
                    add_paragraph(doc, "Figure 20 uses a textual fallback because the gallery HTML contains 19 rendered panels. The report still preserves the required figure slot and caption sequence.", align=WD_ALIGN_PARAGRAPH.JUSTIFY)
            continue
        if line.startswith("[INSERT TABLE HERE]"):
            add_subheading(doc, line.strip())
            idx += 1
            continue
        if re.match(r"^\d+\.\s", line):
            add_paragraph(doc, line, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
            idx += 1
            continue
        if line.startswith("- "):
            add_paragraph(doc, line[2:], align=WD_ALIGN_PARAGRAPH.JUSTIFY)
            idx += 1
            continue
        # Keep the body as thesis-style paragraphs.
        add_paragraph(doc, line, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
        idx += 1


def main() -> None:
    if not SOURCE_PATH.exists():
        raise FileNotFoundError(f"Missing source markdown: {SOURCE_PATH}")

    text = SOURCE_PATH.read_text(encoding="utf-8")
    pages = parse_markdown_sections(text)

    doc = Document()
    configure_styles(doc)
    set_margins(doc)
    set_page_number_footer(doc)

    # Use the markdown's own opening front matter too.
    preface_lines = []
    for line in text.splitlines():
        if line.startswith("=== Page "):
            break
        preface_lines.append(line)
    preface = [line for line in preface_lines if line.strip()]
    if preface:
        add_title(doc, preface[0].lstrip("# "))
        for line in preface[1:]:
            add_paragraph(doc, line, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
        doc.add_page_break()

    for index, (page_num, body, lines) in enumerate(pages):
        render_page(doc, page_num, body, lines)
        if index < len(pages) - 1:
            doc.add_page_break()

    doc.save(str(OUTPUT_PATH))
    print(f"Generated DOCX: {OUTPUT_PATH}")
    print(f"Source pages: {len(pages)}")
    print(f"Size KB: {OUTPUT_PATH.stat().st_size / 1024:.1f}")


if __name__ == "__main__":
    main()
